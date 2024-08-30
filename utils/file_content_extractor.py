from utils.logging_setup import log_file_content_extractor
from config.file_constants import FileConstants
import os
import subprocess
import PyPDF2
from docx import Document
import magic
import patoolib
import shutil
import tempfile
import vobject
import gnupg
from io import BytesIO
from bs4 import BeautifulSoup
from odf.opendocument import load
from odf.text import P
from openpyxl import load_workbook


class FileContentExtractor:
    def __init__(self, file_path=None, content=None):
        self.file_path = file_path
        self.content = content
        self.gpg = gnupg.GPG()

        if not file_path and not content:
            raise ValueError("Either file_path or content must be provided.")
        if file_path and content:
            raise ValueError("Both file_path and content cannot be provided simultaneously.")

        self._check_content_source()

    def _check_content_source(self):
        mime = magic.Magic(mime=True)
        if self.file_path:
            self.file_mime_type = mime.from_file(self.file_path)
        elif self.content:
            self.file_mime_type = mime.from_buffer(self.content)
    def extract_text(self):
        match self.file_mime_type:
            case 'application/pdf':
                return self._extract_text_from_pdf()
            case 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return self._extract_text_from_docx()
            case 'application/msword':
                return self._extract_text_from_doc()
            case 'application/vnd.oasis.opendocument.text':
                return self._extract_text_from_odt()
            case 'application/vnd.oasis.opendocument.spreadsheet':
                # Todo !!!
                pass
            case 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                return self._extract_text_from_excel()
            case 'application/vnd.ms-excel':
                # Todo !!!
                pass
            case 'application/vnd.ms-outlook':
                # Todo !!!
                pass
            case 'text/plain' | 'text/x-shellscript':
                return self._extract_text_from_txt()
            case 'text/html':
                return self._extract_text_from_html()
            case 'text/rtf':
                return self._extract_text_from_rtf()
            case 'text/vcard' | 'text/x-vcard':
                return self._extract_text_from_vcf()
            case 'text/calendar':
                return self._extract_text_from_calendar()
            case 'inode/x-empty':
                print("Empty file detected. No content to extract.")
                return ""
            case mime if 'image' in mime:
                print("Image extension detected")
            case mime if 'video' in mime:
                print("Video extension detected")
            case mime if self._is_archive(mime):
                return self._extract_from_archive()
            case 'application/pgp-keys':
                return self._extract_text_from_pgp_key()
            case 'application/pgp-encrypted':
                return self._decrypt_pgp_encrypted_file()
            case 'application/pgp-signature':
                # Todo !!!
                pass
            case _:
                #raise ValueError(f"Unsupported file type: {self.file_mime_type}")
                log_file_content_extractor.debug(f"Unsupported file type: {self.file_mime_type}")

    def _is_archive(self, mime_type):
        # Les types MIME pour les archives
        archive_mime_types = [
            'application/zip',
            'application/x-rar',
            'application/x-rar-compressed',
            'application/x-7z-compressed',
            'application/gzip',
            'application/x-tar',
            'application/x-bzip2',
        ]
        return mime_type in archive_mime_types

    def _extract_from_archive(self):
        temp_dir = tempfile.mkdtemp()
        try:
            if self.file_path:
                # Extraire les fichiers de l'archive dans un dossier temporaire
                patoolib.extract_archive(self.file_path, outdir=temp_dir)
            elif self.content:
                with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                    temp_file.write(self.content)
                    temp_file.close()
                    patoolib.extract_archive(temp_file.name, outdir=temp_dir)

            extracted_text = []
            # Parcourir les fichiers extraits et les traiter récursivement
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    sub_extractor = FileContentExtractor(file_path=file_path)
                    try:
                        text = sub_extractor.extract_text()
                        if text:
                            extracted_text.append(sub_extractor.extract_text())
                    except ValueError as e:
                        print(f"Unsupported file type {file} inside archive: {e}")

            return "\n".join(extracted_text)
        finally:
            # Nettoyer le dossier temporaire
            shutil.rmtree(temp_dir)

    def _extract_text_from_pdf(self):
        text = ""

        # First attempt using PyPDF2
        try:
            if self.file_path:
                with open(self.file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text()
            elif self.content:
                reader = PyPDF2.PdfReader(BytesIO(self.content))
                for page in reader.pages:
                    text += page.extract_text()
            return text

        except Exception as e:
            print(f"PyPDF2 failed: {e}")

        # Second attempt using PyMuPDF (fitz)
        import fitz
        try:
            if self.file_path:
                doc = fitz.open(self.file_path)
            elif self.content:
                doc = fitz.open(stream=self.content, filetype="pdf")

            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text

        except Exception as e:
            print(f"PyMuPDF failed: {e}")

        # Final attempt using pdfminer.six
        from pdfminer.high_level import extract_text as pdfminer_extract_text
        try:
            if self.file_path:
                text = pdfminer_extract_text(self.file_path)
            elif self.content:
                with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                    temp_file.write(self.content)
                    temp_file.close()
                    text = pdfminer_extract_text(temp_file.name)
                    os.remove(temp_file.name)
            return text

        except Exception as e:
            raise Exception(f"Failed to extract text from PDF using all available methods: {e}")

    def _extract_text_from_docx(self):
        if self.file_path:
            doc = Document(self.file_path)
        elif self.content:
            doc = Document(BytesIO(self.content))

        return "\n".join([para.text for para in doc.paragraphs])

    def _extract_text_from_doc(self):
        if self.file_path:
            output_docx_path = os.path.splitext(self.file_path)[0] + ".docx"
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'docx',
                self.file_path, '--outdir', os.path.dirname(self.file_path)
            ])
            if not os.path.exists(output_docx_path):
                raise FileNotFoundError(f"Conversion failed, {output_docx_path} not found.")
            self.file_path = output_docx_path
            text = self._extract_text_from_docx()
            os.remove(output_docx_path)
            return text
        elif self.content:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                temp_file.write(self.content)
                temp_file.close()
                temp_extractor = FileContentExtractor(file_path=temp_file.name)
                text = temp_extractor._extract_text_from_doc()
                os.remove(temp_file.name)
                return text

    def _extract_text_from_odt(self):
        """
        Extracts text content from an ODT file (OpenDocument Text).
        """
        try:
            doc = load(self.file_path)
            paragraphs = []
            for element in doc.getElementsByType(P):
                paragraphs.append(str(element))

            return "\n".join(paragraphs)
        except Exception as e:
            print(f"Error processing ODT file: {e}")
            return ""

    def _extract_text_from_excel(self):
        text = ""
        if self.file_path:
            workbook = load_workbook(filename=self.file_path, data_only=True)
        else:
            workbook = load_workbook(filename=BytesIO(self.content), data_only=True)

        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                text += row_text + "\n"
        return text

    def _extract_text_from_txt(self):
        if self.file_path:
            with open(self.file_path, 'rb') as file:
                content = file.read()
        elif self.content:
            content = self.content
        text = None
        for encoding in FileConstants.SUPPORTED_ENCODINGS:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                log_file_content_extractor.warning(f"Failed to decode with {encoding}")
        if text is None:
            raise UnicodeDecodeError("Unable to decode text with supported encodings.")
        return text

    def _extract_text_from_html(self):
        """
        Extracts text content from an HTML file (text/html).
        """
        try:
            if self.file_path:
                with open(self.file_path, 'r', encoding='utf-8') as file:
                    html_content = file.read()
            elif self.content:
                html_content = self.content.decode('utf-8')

            soup = BeautifulSoup(html_content, 'html.parser')
            return soup.get_text(separator='\n', strip=True)

        except Exception as e:
            print(f"Error processing HTML file: {e}")
            return ""

    def _extract_text_from_rtf(self):
        from striprtf.striprtf import rtf_to_text
        if self.file_path:
            with open(self.file_path, 'r') as file:
                rtf_content = file.read()
        elif self.content:
            rtf_content = self.content.decode('utf-8')
        return rtf_to_text(rtf_content)

    def _extract_text_from_pgp_key(self):
        if self.file_path:
            with open(self.file_path, 'r') as key_file:
                key_data = key_file.read()
        elif self.content:
            key_data = self.content.decode('utf-8')
        return key_data

    def _decrypt_pgp_encrypted_file(self):
        if self.file_path:
            with open(self.file_path, 'rb') as f:
                decrypted_data = self.gpg.decrypt_file(f)
        elif self.content:
            decrypted_data = self.gpg.decrypt(self.content)

        if not decrypted_data.ok:
            raise ValueError(f"Failed to decrypt PGP file: {decrypted_data.status}")

        # Sauvegarde du contenu décrypté dans un fichier temporaire
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file.write(decrypted_data.data)
        temp_file.close()

        # Récursivement extraire le texte du fichier décrypté
        sub_extractor = FileContentExtractor(temp_file.name)
        extracted_text = sub_extractor.extract_text()

        os.remove(temp_file.name)
        return extracted_text

    def _extract_text_from_vcf(self):
        contacts_info = []

        if self.file_path:
            vcf_source = open(self.file_path, 'r')
        elif self.content:
            vcf_source = BytesIO(self.content).read().decode('utf-8').splitlines()

        vcf_file_iterator = iter(vcf_source)

        while True:
            try:
                vcard = vobject.readOne(vcf_file_iterator)
                contact_info = []
                if vcard.fn.value:
                    contact_info.append(f"Full Name: {vcard.fn.value}")
                if hasattr(vcard, 'email'):
                    contact_info.append(f"Email: {vcard.email.value}")
                if hasattr(vcard, 'tel'):
                    contact_info.append(f"Phone Number: {vcard.tel.value}")
                contacts_info.append("\n".join(contact_info))
            except StopIteration:
                break
            except Exception as e:
                print(f"Error processing vCard: {e}")
                continue

        if self.file_path:
            vcf_source.close()

        return "\n".join(contacts_info)

    def _extract_text_from_calendar(self):
        """
        Extracts text content from a calendar file (text/calendar).
        """
        try:
            if self.file_path:
                with open(self.file_path, 'r') as file:
                    calendar_data = file.read()
            elif self.content:
                calendar_data = self.content.decode('utf-8')

            calendar = vobject.readOne(calendar_data)
            extracted_text = []
            for component in calendar.components():
                for name, prop in component.contents.items():
                    extracted_text.append(f"{name}: {prop[0].value}")

            return "\n".join(extracted_text)

        except Exception as e:
            print(f"Error processing calendar file: {e}")
            return ""


if __name__ == "__main__":
    file_path = "fee82a7202eb32c93010e141a6f601ef710b8343bfc32b0a273e30908c0b254d.doc"
    file_text_extractor = FileContentExtractor(file_path=file_path)
    print(file_text_extractor.extract_text())