import os
import subprocess
import PyPDF2
from docx import Document
import magic
import patoolib
import shutil
import tempfile
import vobject
import gnupg
from io import BytesIO

class FileTextExtractor:
    def __init__(self, file_path=None, content=None):
        self.file_path = file_path
        self.content = content
        self.gpg = gnupg.GPG()

        if not file_path and not content:
            raise ValueError("Either file_path or content must be provided.")
        if file_path and content:
            raise ValueError("Both file_path and content cannot be provided simultaneously.")

        self._check_content_source()

    def _check_content_source(self):
        mime = magic.Magic(mime=True)
        if self.file_path:
            self.file_mime_type = mime.from_file(self.file_path)
        elif self.content:
            self.file_mime_type = mime.from_buffer(self.content)
    def extract_text(self):
        if self.file_mime_type == 'application/pdf':
            return self._extract_text_from_pdf()
        elif self.file_mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self._extract_text_from_docx()
        elif self.file_mime_type == 'application/msword':
            return self._extract_text_from_doc()
        elif self.file_mime_type == 'text/plain':
            return self._extract_text_from_txt()
        elif self.file_mime_type in ['text/vcard', 'text/x-vcard']:
            # ToDo: implementation to be revised !!!
            return self._extract_text_from_vcf()
        elif 'image' in self.file_mime_type:
            print("Image extension detected")
        elif 'video' in self.file_mime_type:
            print("Video extension detected")
        elif self._is_archive(self.file_mime_type):
            return self._extract_from_archive()
        elif self.file_mime_type == 'application/pgp-keys':
            return self._extract_text_from_pgp_key()
        elif self.file_mime_type == 'application/pgp-encrypted':
            return self._decrypt_pgp_encrypted_file()
        else:
            # raise ValueError(f"Unsupported file type: {file_extension}")
            raise ValueError(f"Unsupported file type: {self.file_mime_type}")

    def _is_archive(self, mime_type):
        # Les types MIME pour les archives
        archive_mime_types = [
            'application/zip',
            'application/x-rar-compressed',
            'application/x-7z-compressed',
            'application/gzip',
            'application/x-tar',
            'application/x-bzip2',
        ]
        return mime_type in archive_mime_types

    def _extract_from_archive(self):
        temp_dir = tempfile.mkdtemp()
        try:
            if self.file_path:
                # Extraire les fichiers de l'archive dans un dossier temporaire
                patoolib.extract_archive(self.file_path, outdir=temp_dir)
            elif self.content:
                with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                    temp_file.write(self.content)
                    temp_file.close()
                    patoolib.extract_archive(temp_file.name, outdir=temp_dir)

            extracted_text = []
            # Parcourir les fichiers extraits et les traiter récursivement
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    sub_extractor = FileTextExtractor(file_path=file_path)
                    try:
                        extracted_text.append(sub_extractor.extract_text())
                    except ValueError as e:
                        print(f"Unsupported file type {file} inside archive: {e}")

            return "\n".join(extracted_text)
        finally:
            # Nettoyer le dossier temporaire
            shutil.rmtree(temp_dir)

    def _extract_text_from_pdf(self):
        if self.file_path:
            with open(self.file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
            return text
        elif self.content:
            reader = PyPDF2.PdfReader(BytesIO(self.content))
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text

    def _extract_text_from_docx(self):
        if self.file_path:
            doc = Document(self.file_path)
        elif self.content:
            doc = Document(BytesIO(self.content))

        return "\n".join([para.text for para in doc.paragraphs])

    def _extract_text_from_doc(self):
        if self.file_path:
            output_docx_path = os.path.splitext(self.file_path)[0] + ".docx"
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'docx',
                self.file_path, '--outdir', os.path.dirname(self.file_path)
            ])
            if not os.path.exists(output_docx_path):
                raise FileNotFoundError(f"Conversion failed, {output_docx_path} not found.")
            self.file_path = output_docx_path
            text = self._extract_text_from_docx()
            os.remove(output_docx_path)
            return text
        elif self.content:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                temp_file.write(self.content)
                temp_file.close()
                temp_extractor = FileTextExtractor(file_path=temp_file.name)
                text = temp_extractor._extract_text_from_doc()
                os.remove(temp_file.name)
                return text

    def _extract_text_from_txt(self):
        if self.file_path:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        elif self.content:
            text = self.content.decode('utf-8')
        return text

    def _extract_text_from_pgp_key(self):
        if self.file_path:
            with open(self.file_path, 'r') as key_file:
                key_data = key_file.read()
        elif self.content:
            key_data = self.content.decode('utf-8')
        return key_data

    def _decrypt_pgp_encrypted_file(self):
        if self.file_path:
            with open(self.file_path, 'rb') as f:
                decrypted_data = self.gpg.decrypt_file(f)
        elif self.content:
            decrypted_data = self.gpg.decrypt(self.content)

        if not decrypted_data.ok:
            raise ValueError(f"Failed to decrypt PGP file: {decrypted_data.status}")

        # Sauvegarde du contenu décrypté dans un fichier temporaire
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file.write(decrypted_data.data)
        temp_file.close()

        # Récursivement extraire le texte du fichier décrypté
        sub_extractor = FileTextExtractor(temp_file.name)
        extracted_text = sub_extractor.extract_text()

        os.remove(temp_file.name)
        return extracted_text

    def _extract_text_from_vcf(self):
        contacts_info = []

        if self.file_path:
            vcf_source = open(self.file_path, 'r')
        elif self.content:
            vcf_source = BytesIO(self.content).read().decode('utf-8').splitlines()

        vcf_file_iterator = iter(vcf_source)

        while True:
            try:
                vcard = vobject.readOne(vcf_file_iterator)
                contact_info = []
                if vcard.fn.value:
                    contact_info.append(f"Full Name: {vcard.fn.value}")
                if hasattr(vcard, 'email'):
                    contact_info.append(f"Email: {vcard.email.value}")
                if hasattr(vcard, 'tel'):
                    contact_info.append(f"Phone Number: {vcard.tel.value}")
                contacts_info.append("\n".join(contact_info))
            except StopIteration:
                break
            except Exception as e:
                print(f"Error processing vCard: {e}")
                continue

        if self.file_path:
            vcf_source.close()

        return "\n".join(contacts_info)


if __name__ == "__main__":
    file_path = "/home/valentin/github/MessageTrackPro_v2/attachments/fee82a7202eb32c93010e141a6f601ef710b8343bfc32b0a273e30908c0b254d.doc"
    file_text_extractor = FileTextExtractor(file_path=file_path)
    print(file_text_extractor.extract_text())